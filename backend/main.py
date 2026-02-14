import time
import random
import csv
import os
import datetime
from fastapi import FastAPI, UploadFile, File, HTTPException, Form
from fastapi.middleware.cors import CORSMiddleware
from pydantic import BaseModel
import uvicorn
import json
from typing import List
from typing import List, Optional
import shutil  # 👈 新增
from fastapi.staticfiles import StaticFiles
from fastapi.responses import FileResponse, RedirectResponse, JSONResponse
from openai import OpenAI
from io import BytesIO
import io
import traceback

# 文件解析库（按需导入，避免依赖问题）
try:
    import PyPDF2
    PDF_SUPPORT = True
except ImportError:
    PDF_SUPPORT = False
    print("⚠️ PyPDF2 未安装，PDF 文件解析功能将不可用")

try:
    from docx import Document
    DOCX_SUPPORT = True
except ImportError:
    DOCX_SUPPORT = False
    print("⚠️ python-docx 未安装，DOCX 文件解析功能将不可用")

# ==========================================
# 导入数据库配置和操作函数
# ==========================================
from .db_config import (
    get_db_connection,
    get_all_users,
    get_user_by_username,
    user_login,
    update_user_field,
    update_user_multiple_fields,
    create_user,
    increment_user_field,
    decrement_user_field,
)

app = FastAPI()

os.makedirs("static/avatars", exist_ok=True)
app.mount("/static", StaticFiles(directory="static"), name="static")

frontend_dist = os.path.join(os.path.dirname(__file__), "..", "frontend", "dist")
if os.path.exists(frontend_dist):
    app.mount("/assets", StaticFiles(directory=os.path.join(frontend_dist, "assets")), name="frontend_assets")

# --- 1. 跨域配置：按需开放，确保支持 OPTIONS/POST/GET 等所有方法 ---
app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],  # 包含 OPTIONS / POST / GET
    allow_headers=["*"],
)

# ==========================================
#  DeepSeek 客户端 (新增：虚拟实验 & 生涯规划整合)
# ==========================================
DEEPSEEK_API_KEY = os.getenv("DEEPSEEK_API_KEY", "sk-d3a066f75e744cd58708b9af635d3606")
deepseek_client = OpenAI(
    api_key=DEEPSEEK_API_KEY,
    base_url="https://api.deepseek.com"
)

def _deepseek_markdown(system_prompt: str, user_prompt: str) -> str:
    """调用 DeepSeek，返回 Markdown 文本"""
    try:
        resp = deepseek_client.chat.completions.create(
            model="deepseek-chat",
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_prompt},
            ],
            temperature=0.4,
        )
        return (resp.choices[0].message.content or "").strip()
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"DeepSeek 调用失败: {e}")

def _deepseek_json(system_prompt: str, user_prompt: str) -> dict:
    """调用 DeepSeek，要求其返回严格 JSON 对象"""
    try:
        resp = deepseek_client.chat.completions.create(
            model="deepseek-chat",
            messages=[
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": user_prompt},
            ],
            temperature=0.3,
            response_format={"type": "json_object"},
        )
        content = resp.choices[0].message.content or "{}"
        return json.loads(content)
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"DeepSeek(JSON) 调用失败: {e}")


def extract_text_from_file(upload_file: UploadFile) -> str:
    """
    从上传的文件中提取文本内容，支持 PDF / DOCX / TXT。
    解析失败时抛出带有明确信息的 HTTPException，避免静默返回空内容。
    """
    try:
        from PyPDF2 import PdfReader
        from docx import Document
    except ImportError as e:
        # 依赖缺失时直接给出明确提示，方便在 Render 等环境排查
        raise HTTPException(status_code=500, detail=f"服务器缺少文件解析依赖，请安装 PyPDF2 和 python-docx: {e}")

    try:
        file_content = upload_file.file.read()
        file_name = (upload_file.filename or "").lower()

        # 基本类型检查
        if file_name.endswith(".pdf"):
            try:
                reader = PdfReader(BytesIO(file_content))
                text_parts = []
                for page in reader.pages:
                    page_text = page.extract_text() or ""
                    text_parts.append(page_text)
                text = "\n".join(text_parts)
            except Exception as e:
                raise HTTPException(status_code=500, detail=f"PDF 解析失败: {str(e)}")

        elif file_name.endswith(".docx"):
            try:
                doc = Document(BytesIO(file_content))
                text = "\n".join([para.text for para in doc.paragraphs])
            except Exception as e:
                raise HTTPException(status_code=500, detail=f"DOCX 解析失败: {str(e)}")

        elif file_name.endswith(".txt"):
            try:
                text = file_content.decode("utf-8")
            except Exception as e:
                raise HTTPException(status_code=500, detail=f"TXT 解码失败: {str(e)}")

        else:
            raise HTTPException(status_code=400, detail="不支持的文件格式，请上传 PDF/DOCX/TXT 文件。")

        if not text or not text.strip():
            raise HTTPException(status_code=400, detail="文件内容为空，请检查文件。")

        return text
    except HTTPException:
        # 保持 HTTPException 语义，直接向上抛给接口处理
        raise
    except Exception as e:
        # 兜底异常处理，确保不会静默失败
        raise HTTPException(status_code=500, detail=f"文件解析失败: {str(e)}")
    finally:
        # 重置文件指针，避免影响后续操作
        try:
            upload_file.file.seek(0)
        except Exception:
            pass

# ==========================================
#  简历文件解析函数
# ==========================================
def extract_text_from_file(upload_file: UploadFile) -> str:
    """从上传的文件中提取文本内容"""
    import PyPDF2
    from docx import Document
    from io import BytesIO

    try:
        # 读取文件内容
        file_content = upload_file.file.read()
        file_name = upload_file.filename.lower() if upload_file.filename else ""

        # 根据文件类型解析内容
        if file_name.endswith(".pdf"):
            reader = PyPDF2.PdfReader(BytesIO(file_content))
            text_parts = []
            for page in reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)
            text = "\n".join(text_parts)
        elif file_name.endswith(".docx"):
            doc = Document(BytesIO(file_content))
            text = "\n".join([para.text for para in doc.paragraphs if para.text.strip()])
        elif file_name.endswith(".txt"):
            try:
                text = file_content.decode("utf-8")
            except UnicodeDecodeError:
                text = file_content.decode("gbk")  # 兼容中文编码
        else:
            raise HTTPException(status_code=400, detail="不支持的文件格式，请上传 PDF/DOCX/TXT 文件。")

        # 检查解析后的内容是否为空
        if not text or not text.strip():
            raise HTTPException(status_code=400, detail="文件内容为空，或无法从文件中提取有效文本。")

        return text.strip()

    except HTTPException:
        # 直接向上抛出，让 /api/analyze_resume 返回对应的 400 提示
        raise
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"文件解析失败: {str(e)}")
    finally:
        # 重置文件指针，以便后续操作
        try:
            upload_file.file.seek(0)
        except Exception:
            pass

# ==========================================
#  模型定义 (整合了所有功能的数据结构)
# ==========================================
class LoginRequest(BaseModel):
    username: str
    password: str

class RegisterRequest(BaseModel):
    username: str
    password: str
    grade: str
    target_role: str

class ChatRequest(BaseModel):
    message: str

class RoadmapRequest(BaseModel):
    current_grade: str
    target_role: str

class AgentRequest(BaseModel):
    grade: str
    target_job: str

class ApplyRequest(BaseModel):
    username: str
    job_name: str
    salary: str

class AdminProfileModel(BaseModel):
    username: str = "admin"
    nickname: str = ""
    role: str = ""
    department: str = ""
    email: str = ""
    phone: str = ""
    avatar: str = ""     # 存 Base64 字符串
    lastLogin: str = ""
    ip: str = ""
    new_password: str = None # 接收新密码（已废弃，使用专门的密码修改接口）

class AdminChangePasswordRequest(BaseModel):
    username: str
    old_password: str
    new_password: str

class GenerateResumeRequest(BaseModel):
    focus_direction: str = "通用"
    diagnosis: dict | None = None

# ==========================================
#  新增功能 G: 虚拟实验体验 & 生涯分析整合
# ==========================================
class AnalyzeExperimentRequest(BaseModel):
    answers: dict
    career: str | None = None

class GenerateCareerRequest(BaseModel):
    personality_json: dict
    experiment_markdown: str
    note: str | None = ""

class VirtualCareerQuestionsRequest(BaseModel):
    career: str

class GenerateJobTestRequest(BaseModel):
    """生成职业虚拟体验 + 15 道测试题的请求体"""
    jobName: str

class GenerateInterviewReportRequest(BaseModel):
    chat_history: list  # 完整的对话历史记录
    target_role: str = ""  # 目标岗位

class AnalyzeCompetitivenessRequest(BaseModel):
    """竞争力分析请求体"""
    gpa: str  # GPA（可以是数字或文字描述）
    project_experience: str  # 项目实战经验
    internship: str  # 名企实习经历
    competition: str  # 竞赛获奖情况
    english_academic: str  # 英语学术能力
    leadership: str  # 领导力与协作
    meta: dict | None = None  # 前端已提取的元信息：身份/方向/时长/生成时间


class AnalyzeNaturalLanguageRequest(BaseModel):
    """竞争力沙盘自然语言量化请求体"""
    text: str

# ==========================================
#  Mock 数据库 (职位数据)
# ==========================================
JOB_DATABASE = [
    {"职业分类": "后端开发", "岗位": "Python 开发工程师", "关键词": "FastAPI, MySQL", "平均薪资": "15k-25k"},
    {"职业分类": "前端开发", "岗位": "Vue 开发工程师", "关键词": "Vue3, Element Plus", "平均薪资": "14k-23k"},
    {"职业分类": "算法工程师", "岗位": "NLP 算法工程师", "关键词": "LLM, RAG", "平均薪资": "20k-35k"},
    {"职业分类": "数据开发", "岗位": "大数据开发工程师", "关键词": "Hadoop, Spark", "平均薪资": "18k-30k"},
    {"职业分类": "测试", "岗位": "自动化测试工程师", "关键词": "Selenium, PyTest", "平均薪资": "12k-20k"},
]

# --- 1. 定义历史记录的数据模型 ---
class HistoryItem(BaseModel):
    username: str
    action_type: str  # "诊断" 或 "生成"
    title: str        # 例如 "Java工程师简历诊断"
    score: int
    date: str
    status: str       # "已完成"
# 1. 获取管理员信息 (GET)
# ⚠️ 之前报错 404 就是因为这个函数可能没写对，或者缩进错了
@app.get("/api/admin/profile")
def get_admin_profile():
    print("🔍 [DEBUG] 收到获取 Admin Profile 请求") # 调试日志
    
    file_path = "data/admin_profile.json"
    
    # 确保 data 目录存在
    if not os.path.exists("data"):
        os.makedirs("data")
    
    # 从数据库读取 admin 用户信息
    db_user = get_user_by_username("admin")
    db_data = {}
    if db_user:
        db_data = {
            "nickname": db_user.get("nickname", ""),
            "phone": db_user.get("phone", ""),
            "email": db_user.get("email", ""),
            "department": db_user.get("department", ""),
        }
    
    # 如果JSON文件不存在，返回默认数据（合并数据库数据）
    if not os.path.exists(file_path):
        print("⚠️ [DEBUG] JSON 文件不存在，返回默认值")
        default_data = {
            "username": "admin",
            "nickname": db_data.get("nickname") or "默认管理员",
            "role": "Super Admin",
            "department": db_data.get("department") or "技术部",
            "email": db_data.get("email") or "admin@careerfly.com",
            "phone": db_data.get("phone") or "13800000000",
            "avatar": ""
        }
        # 写入文件
        with open(file_path, "w", encoding="utf-8") as f:
            json.dump(default_data, f, ensure_ascii=False, indent=2)
        return {"success": True, "data": default_data}
    
    # 读取JSON文件（包含头像等完整信息）
    try:
        with open(file_path, "r", encoding="utf-8") as f:
            data = json.load(f)
        
        # 如果数据库中有更新的数据，合并到返回结果（但保留JSON中的头像）
        if db_data:
            data.update({k: v for k, v in db_data.items() if v})  # 只更新非空字段
        
        print("✅ [DEBUG] 成功读取 JSON 数据")
        return {"success": True, "data": data}
    except Exception as e:
        print(f"❌ [DEBUG] 读取失败: {e}")
        return {"success": False, "message": "读取失败"}
# 2. 更新管理员信息 (POST) - 同步到 JSON 和 CSV
@app.post("/api/admin/profile/update")
def update_admin_profile(item: AdminProfileModel):
    print(f"📝 [DEBUG] 收到更新请求: 昵称={item.nickname}, 头像长度={len(item.avatar) if item.avatar else 0}")

    # --- A. 保存到 JSON (头像、昵称等基本信息) ---
    json_path = "data/admin_profile.json"
    try:
        # 使用 model_dump 替代 dict (修复 Pydantic 警告)
        save_data = item.model_dump(exclude={"new_password"}) 
        
        # 检查头像Base64字符串长度（200KB图片转Base64后约270KB）
        if item.avatar and len(item.avatar) > 300000:  # 约300KB的Base64字符串
            return {"success": False, "message": "头像文件过大，请上传小于200KB的图片"}
        
        with open(json_path, "w", encoding="utf-8") as f:
            json.dump(save_data, f, ensure_ascii=False, indent=2)
        print("✅ [DEBUG] JSON 文件保存成功")
    except Exception as e:
        print(f"❌ [DEBUG] JSON 保存失败: {e}")
        return {"success": False, "message": f"保存失败: {e}"}

    # --- B. 同步所有字段到数据库 (昵称、手机号、邮箱、部门/职位) ---
    try:
        # 准备要更新的字段（不包含头像，头像只存在JSON中）
        update_fields = {}
        if item.nickname:
            update_fields["nickname"] = item.nickname
        if item.phone:
            update_fields["phone"] = item.phone
        if item.email:
            update_fields["email"] = item.email
        if item.department:
            update_fields["department"] = item.department
        
        # 更新数据库
        if update_fields:
            username = item.username or "admin"
            success = update_user_multiple_fields(username, update_fields)
            if success:
                print("✅ [DEBUG] 数据库资料同步完成")
            else:
                print("⚠️ [DEBUG] 数据库更新失败或用户不存在")
    except Exception as e:
        print(f"❌ [DEBUG] 数据库操作出错: {e}")
        # 数据库同步失败不影响JSON保存
        return {"success": True, "message": "资料已保存到JSON，但数据库同步失败"}

    return {"success": True, "message": "更新成功"}

# 3. 管理员密码修改接口（包含旧密码验证）
@app.post("/api/admin/profile/change-password")
def change_admin_password(req: AdminChangePasswordRequest):
    print(f"🔐 [DEBUG] 收到密码修改请求: 用户={req.username}")
    
    # 1. 验证旧密码（使用数据库）
    user = get_user_by_username(req.username)
    if not user:
        return {"success": False, "message": "用户不存在"}
    
    # 验证旧密码
    if user.get("password", "").strip() != req.old_password:
        print(f"❌ [DEBUG] 旧密码不正确")
        return {"success": False, "message": "旧密码不正确，请重新输入"}
    
    # 2. 新密码复杂度校验
    if len(req.new_password) < 8:
        return {"success": False, "message": "新密码长度至少 8 位"}
    
    # 3. 更新数据库密码
    try:
        success = update_user_field(req.username, "password", req.new_password)
        if success:
            print("✅ [DEBUG] 密码更新成功")
            return {"success": True, "message": "密码修改成功，请重新登录"}
        else:
            return {"success": False, "message": "密码更新失败"}
    except Exception as e:
        print(f"❌ [DEBUG] 密码更新失败: {e}")
        return {"success": False, "message": f"密码更新失败: {e}"}

# --- 3. 新增：获取历史记录接口 ---
@app.get("/api/history")
def get_history(username: str):
    file_path = "data/history.csv"
    if not os.path.exists(file_path):
        return {"success": True, "data": []}
    
    records = []
    with open(file_path, "r", encoding="utf-8-sig") as f:
        reader = csv.DictReader(f)
        for row in reader:
            if row['username'] == username:
                records.append(row)
    
    # 按时间倒序排列 (最新的在前面)
    records.reverse()
    return {"success": True, "data": records}
# 根路径处理（避免重复声明 / 路由）
@app.get("/")
def root():
    return {"ok": True, "service": "ai-career-helper-backend"}

@app.get("/health")
def health():
    """健康检查接口"""
    result = {"ok": True}
    
    # 可选的数据库连接检查
    try:
        from .db_config import get_db_connection_with_error
        conn, db_error = get_db_connection_with_error()
        if conn:
            conn.close()
            result["db_ok"] = True
        else:
            result["db_ok"] = False
            if db_error:
                result["db_error"] = db_error
    except Exception as e:
        result["db_ok"] = False
        # 确保错误信息不包含敏感信息（如密码）
        error_msg = str(e)
        # 移除可能包含密码的错误信息
        if "password" in error_msg.lower() or "pwd" in error_msg.lower():
            result["db_error"] = "数据库连接失败"
        else:
            result["db_error"] = error_msg
    
    return result

# 简历医生服务地址配置
RESUME_DOCTOR_URL = os.getenv(
    "RESUME_DOCTOR_URL",
    "https://ai-career-apper-resume-doctor-69etycfa4ohbkxndweoawk.streamlit.app"
)

@app.get("/resume-doctor")
async def redirect_resume_doctor():
    """简历医生服务代理接口"""
    return RedirectResponse(url=RESUME_DOCTOR_URL)

@app.post("/api/login")
def login(request: LoginRequest):
    # 使用数据库验证登录
    success, message = user_login(request.username, request.password)
    if success:
        # 登录成功，获取用户完整信息
        user = get_user_by_username(request.username)
        if user:
            return {
                "success": True, 
                "message": "登录成功", 
                "user": user
            }
        else:
            return {"success": False, "message": "获取用户信息失败"}
    else:
        return {"success": False, "message": message}

# ==========================================
# 🛑 替换 main.py 里的 register 函数
# ==========================================

@app.post("/api/register")
def register(req: RegisterRequest):
    # 使用数据库创建新用户
    success, message = create_user(req.username, req.password, req.grade, req.target_role)
    return {"success": success, "message": message}

# ==========================================
#  核心功能 B: 职位推荐 (修复 404 错误)
# ==========================================
@app.post("/api/recommend")
def recommend():
    """简单的职位推荐接口"""
    return {"success": True, "data": JOB_DATABASE}

# ==========================================
#  核心功能 C: AI 模拟面试 (聊天)
# ==========================================
@app.post("/api/chat")
def chat(request: ChatRequest):
    time.sleep(0.5)
    followups = [
        "在这个项目中，你遇到的最大技术难点是什么？",
        "如果让你优化数据库查询，你会怎么做？",
        "对于高并发场景，你有什么设计思路？"
    ]
    return {
        "reply": f"收到！关于'{request.message}'，我的看法是... (模拟AI回复)\n\n👉 追问：{random.choice(followups)}",
        "meta": {"topic": "技术", "difficulty": "中等"}
    }

# ==========================================
#  核心功能 D: 生涯规划 (雷达图 + 时间轴)
# ==========================================
def _generate_roadmap_with_ai(current_grade: str, target_role: str) -> dict:
    """
    使用 Deepseek API 生成生涯路径规划
    返回结构化的规划数据，包含从大一到大四的详细时间规划表
    """
    system_prompt = """你是一位资深的大学生职业规划导师，拥有10年以上行业经验，擅长为不同年级和职业方向的学生制定详细的成长路径规划。

请根据用户输入的「当前年级」和「意向职业」，生成一份从大一到大四的完整生涯规划。

**输出要求（严格 JSON 格式）：**
{
  "stages": [
    {
      "grade": "大一",
      "title": "阶段标题（如：夯实编程基础）",
      "content": "详细的学习重点、比赛推荐、实习安排和目标荣誉（150-200字，要求具体、可执行）",
      "resources": ["推荐资源1（具体书名或平台名）", "推荐资源2", "推荐资源3"],
      "certificates": ["目标证书/荣誉1（具体名称）", "目标证书/荣誉2"],
      "recommended_companies": []  // 仅大四阶段填写，其他阶段为空数组
    },
    // ... 大二、大三、大四（必须包含所有四个年级）
  ],
  "ai_comment": "针对性的AI导师洞察（80-120字，包含当前阶段重点建议）"
}

**规划原则（必须严格遵守）：**
- **大一阶段**：侧重基础学习与入门竞赛
  * 学习重点：编程语言基础（如Python/Java/C++）、数据结构与算法、计算机基础课程
  * 比赛推荐：蓝桥杯、PAT、CCF CSP认证、校内编程竞赛
  * 实习安排：暂不安排（大一以学习为主）
  * 目标荣誉：计算机二级证书、编程语言认证、校内竞赛获奖、优秀学生奖学金

- **大二阶段**：侧重竞赛进阶与项目实践
  * 学习重点：深入学习专业课程、算法进阶、框架学习、项目开发
  * 比赛推荐：ACM-ICPC、蓝桥杯国赛、天池/Kaggle数据竞赛、开源项目贡献
  * 实习安排：可尝试暑期短期实习或项目实践
  * 目标荣誉：算法竞赛获奖证书、Kaggle竞赛证书、GitHub项目认证、技术博客认证

- **大三阶段**：侧重实习积累与技术深度
  * 学习重点：系统设计、技术深度、行业认知、软技能提升
  * 比赛推荐：继续参与高级竞赛、技术挑战赛、创新项目大赛
  * 实习安排：投递大厂日常实习或暑期实习（如字节跳动、腾讯、阿里巴巴等）
  * 目标荣誉：大厂实习证明、论文发表/技术博客认证、开源项目贡献认证、优秀实习生

- **大四阶段**：侧重校招冲刺与入职准备
  * 学习重点：面试准备、系统设计、简历优化、作品集完善
  * 比赛推荐：继续参与高级竞赛（如时间允许）
  * 实习安排：可继续实习或准备秋招
  * 目标荣誉：算法工程师Offer、技术专家认证、优秀毕业生
  * 推荐企业：根据职业方向推荐3-6家适配企业，并附带简要匹配理由

**企业推荐要求（仅大四阶段）：**
- 根据职业方向推荐3-6家适配的企业
- 包含互联网大厂、知名企业等（如：字节跳动、阿里巴巴、腾讯、百度、美团、滴滴、京东、网易等）
- 每个企业名称简洁（2-4个字）
- 企业选择要贴合职业方向（如算法方向推荐AI/算法强项的企业，前端方向推荐前端技术栈先进的企业）

**内容要求：**
- content 字段必须详细具体，包含具体的学习内容、比赛名称、实习建议、荣誉目标
- resources 字段推荐3-5个具体的学习资源（如书名、在线课程、平台名称）
- certificates 字段推荐2-4个具体可获得的证书或荣誉
- 所有内容必须贴合职业方向，具有针对性和可执行性

请确保输出为严格有效的 JSON 格式，不要包含任何 Markdown 代码块标记（如```json```）。"""

    user_prompt = f"""当前年级：{current_grade}
意向职业：{target_role}

请为这位学生生成一份从大一到大四的详细生涯路径规划。

**要求：**
1. 必须生成完整的四个年级（大一、大二、大三、大四）的规划
2. 每个阶段的内容要详细具体，包含：
   - 具体的学习重点（如：学习哪些技术栈、完成哪些项目）
   - 具体的比赛推荐（如：蓝桥杯、ACM-ICPC、Kaggle等）
   - 具体的实习安排建议（大三、大四阶段）
   - 具体的目标荣誉/证书（如：计算机二级、算法竞赛获奖等）
3. 大四阶段必须包含3-6家适配入职企业的推荐
4. 所有内容必须贴合「{target_role}」这个职业方向，具有针对性和可执行性
5. 考虑到用户当前是「{current_grade}」，请在ai_comment中给出当前阶段的重点建议

请生成规划内容。"""

    try:
        # 调用 Deepseek API 生成规划内容
        ai_response = _deepseek_json(system_prompt, user_prompt)
        
        # 解析 AI 返回的数据
        stages_data = ai_response.get("stages", [])
        ai_comment = ai_response.get("ai_comment", f"基于{current_grade}和{target_role}方向，为你规划了从大一到大四的完整成长路径。")
        
        # 验证数据完整性
        if not stages_data or len(stages_data) < 4:
            print(f"⚠️ AI 返回的阶段数据不完整，共 {len(stages_data)} 个阶段")
        
        # 定义年级索引（用于判断当前进度）
        grade_index = {"大一": 0, "大二": 1, "大三": 2, "大四": 3}.get(current_grade, 1)
        grade_list = ["大一", "大二", "大三", "大四"]
        
        # 构建里程碑数据，确保格式与现有结构完全一致
        stages = []
        for idx, grade in enumerate(grade_list):
            # 从 AI 返回的数据中查找对应年级的规划
            stage_data = None
            for s in stages_data:
                if s.get("grade") == grade:
                    stage_data = s
                    break
            
            # 如果 AI 没有返回该年级的数据，使用默认值（降级处理）
            if not stage_data:
                print(f"⚠️ AI 未返回 {grade} 阶段数据，使用默认模板")
                stage_data = {
                    "title": f"{grade}阶段规划",
                    "content": f"根据{target_role}方向，制定{grade}阶段的学习和实践计划。建议重点关注专业课程学习、项目实践和技能提升。",
                    "resources": ["相关学习资源", "在线课程平台", "技术社区"],
                    "certificates": ["相关证书"],
                    "recommended_companies": []
                }
            
            # 确保必要字段存在
            if not stage_data.get("title"):
                stage_data["title"] = f"{grade}阶段规划"
            if not stage_data.get("content"):
                stage_data["content"] = f"根据{target_role}方向，制定{grade}阶段的学习和实践计划。"
            if not stage_data.get("resources"):
                stage_data["resources"] = ["相关学习资源"]
            if not stage_data.get("certificates"):
                stage_data["certificates"] = ["相关证书"]
            if grade != "大四":
                stage_data["recommended_companies"] = []
            
            # 判断状态：已完成、进行中、等待中
            if idx < grade_index:
                status = "done"
                color = "#67C23A"  # 绿色
                icon = "CircleCheck"
            elif idx == grade_index:
                status = "process"
                color = "#409EFF"  # 蓝色
                icon = "Loading"
            else:
                status = "wait"
                color = "#909399"  # 灰色
                icon = ""
            
            # 构建里程碑数据（格式与现有结构完全一致）
            milestone = {
                "time": grade,
                "title": stage_data.get("title", f"{grade}阶段规划"),
                "content": stage_data.get("content", ""),
                "status": status,
                "color": color,
                "icon": icon,
                "resources": stage_data.get("resources", []),
                "certificates": stage_data.get("certificates", []),
                "timestamp": f"{grade}学年"
            }
            
            # 如果是大四阶段，添加推荐企业
            if grade == "大四" and stage_data.get("recommended_companies"):
                milestone["recommended_companies"] = stage_data.get("recommended_companies", [])
            
            stages.append(milestone)
        
        return {
            "stages": stages,
            "ai_comment": ai_comment
        }
        
    except Exception as e:
        print(f"❌ AI 生成生涯规划失败: {e}")
        # 如果 AI 生成失败，返回默认规划（降级处理）
        raise HTTPException(status_code=500, detail=f"AI 生成生涯规划失败: {str(e)}")

@app.post("/api/generate_roadmap")
def generate_roadmap(req: RoadmapRequest):
    """
    AI 驱动的生涯路径规划生成
    使用 Deepseek API 基于用户输入的年级和意向方向，生成个性化的关键里程碑规划
    """
    # 雷达图逻辑（保持不变）
    radar_indicators = [
        {"name": "基础知识", "max": 100}, {"name": "实战能力", "max": 100},
        {"name": "算法思维", "max": 100}, {"name": "工程素养", "max": 100},
        {"name": "软技能", "max": 100}
    ]
    base_score = 60 if "大一" in req.current_grade else (70 if "大二" in req.current_grade else 80)
    current_scores = [base_score + random.randint(-10, 10) for _ in range(5)]

    # ==========================================
    # 使用 AI 生成关键里程碑规划（替换原有固定模板）
    # ==========================================
    try:
        ai_result = _generate_roadmap_with_ai(req.current_grade, req.target_role)
        stages = ai_result["stages"]
        ai_comment = ai_result["ai_comment"]
    except Exception as e:
        # 如果 AI 生成失败，使用降级方案（保留原有模板逻辑作为备选）
        print(f"⚠️ AI 生成失败，使用降级方案: {e}")
        grade = req.current_grade
        direction = req.target_role
        grade_index = {"大一": 0, "大二": 1, "大三": 2, "大四": 3}.get(grade, 1)
        
        # 降级方案：使用简化的固定模板（仅作为备选，正常情况下不会执行）
        direction_templates = {
            "算法": {
                "大一": {
                    "title": "夯实编程基础",
                    "content": "系统学习 C++/Python 基础语法，完成数据结构与算法课程，开始刷 LeetCode（目标：100题）",
                    "resources": ["《算法导论》", "LeetCode 刷题计划", "Python 基础教程"],
                    "certificates": ["计算机二级证书", "Python 编程认证"]
                },
                "大二": {
                    "title": "算法竞赛与深度学习入门",
                    "content": "参加 ACM/蓝桥杯等算法竞赛，学习机器学习基础（线性代数、概率论），完成第一个深度学习项目（如手写数字识别）",
                    "resources": ["《机器学习》- 周志华", "Kaggle 竞赛", "PyTorch 官方教程"],
                    "certificates": ["算法竞赛获奖证书", "Kaggle 竞赛证书"]
                },
                "大三": {
                    "title": "算法实习与项目实战",
                    "content": "投递算法实习岗位（如字节跳动、腾讯 AI Lab），参与 NLP/CV 相关项目，发表技术博客或论文",
                    "resources": ["《深度学习》- Ian Goodfellow", "GitHub 开源项目", "技术博客平台"],
                    "certificates": ["大厂算法实习证明", "论文发表/技术博客认证"]
                },
                "大四": {
                    "title": "秋招冲刺与职业定位",
                    "content": "准备算法工程师秋招（刷题 300+，准备系统设计），完善简历和作品集，目标企业：字节、阿里、腾讯、百度等",
                    "resources": ["《剑指 Offer》", "系统设计面试指南", "算法面试真题"],
                    "certificates": ["算法工程师 Offer", "技术专家认证"],
                    "recommended_companies": ["字节跳动", "阿里巴巴", "腾讯", "百度", "美团", "滴滴"]
                }
            }
        }
        
        # 获取对应方向的模板，如果没有则使用算法模板
        template = direction_templates.get(direction, direction_templates.get("算法", {}))
        
        # 生成四个阶段的里程碑
        stages = []
        grade_list = ["大一", "大二", "大三", "大四"]
        
        for idx, g in enumerate(grade_list):
            stage_data = template.get(g, {
                "title": f"{g}阶段规划",
                "content": f"根据{direction}方向，制定{g}阶段的学习和实践计划",
                "resources": ["相关学习资源"],
                "certificates": ["相关证书"]
            })
            
            # 判断状态：已完成、进行中、等待中
            if idx < grade_index:
                status = "done"
                color = "#67C23A"  # 绿色
                icon = "CircleCheck"
            elif idx == grade_index:
                status = "process"
                color = "#409EFF"  # 蓝色
                icon = "Loading"
            else:
                status = "wait"
                color = "#909399"  # 灰色
                icon = ""
            
            # 构建里程碑数据
            milestone = {
                "time": g,
                "title": stage_data["title"],
                "content": stage_data["content"],
                "status": status,
                "color": color,
                "icon": icon,
                "resources": stage_data.get("resources", []),
                "certificates": stage_data.get("certificates", []),
                "timestamp": f"{g}学年"
            }
            
            # 如果是大四阶段，添加推荐企业
            if g == "大四" and "recommended_companies" in stage_data:
                milestone["recommended_companies"] = stage_data["recommended_companies"]
            
            stages.append(milestone)
        
        # 生成 AI 评论
        ai_comment = f"基于{grade}和{direction}方向，为你规划了从大一到大四的完整成长路径。当前处于{grade}阶段，建议重点关注{stages[grade_index]['title']}，为下一阶段做好准备。"
    
    # 返回统一格式的数据（无论 AI 生成成功还是降级方案）
    return {
        "radar_chart": {"indicators": radar_indicators, "values": current_scores},
        "ai_comment": ai_comment,
        "roadmap": stages
    }

# ==========================================
#  核心功能 E: Agent 职位推荐 & 投递
# ==========================================
@app.post("/api/agent")
def agent_recommend(req: AgentRequest):
    recommendations = [j for j in JOB_DATABASE if req.target_job in j['岗位'] or req.target_job in j['职业分类']]
    if not recommendations: recommendations = JOB_DATABASE[:2]
    
    return {
        "reply": f"我是你的 Agent。为你找到 {len(recommendations)} 个相关岗位。",
        "data": recommendations
    }

@app.post("/api/apply")
def apply_job(req: ApplyRequest):
    file_path = "data/applications.csv"
    os.makedirs("data", exist_ok=True)
    if not os.path.exists(file_path):
        with open(file_path, "w", encoding="utf-8-sig", newline="") as f:
            csv.writer(f).writerow(["用户", "岗位", "薪资", "时间", "状态"])
    
    with open(file_path, "a", encoding="utf-8-sig", newline="") as f:
        csv.writer(f).writerow([req.username, req.job_name, req.salary, datetime.datetime.now(), "已投递"])
    return {"message": "投递成功", "status": "success"}

# ... 之前的代码 ...

# --- 1. 定义用户资料模型 ---
class UserProfile(BaseModel):
    username: str
    avatar: str = ""  # 👈 新增这一行
    email: str = ""
    phone: str = ""
    city: str = ""
    style: str = "专业正式"
    file_format: str = "PDF"
    notify: bool = True
    auto_save: bool = True

# --- 2. 获取用户资料接口 ---
@app.get("/api/user/profile")
def get_profile(username: str):
    """
    获取用户资料接口
    
    优先从数据库加载，如果数据库没有则从CSV加载
    确保数据持久化，刷新后不丢失
    """
    # 1. 优先从数据库加载
    try:
        user = get_user_by_username(username)
        if user:
            # 从数据库获取用户信息
            profile_data = {
                "username": username,
                "email": getattr(user, 'email', '') or '',
                "phone": getattr(user, 'phone', '') or '',
                "city": getattr(user, 'city', '') or '',
                "avatar": getattr(user, 'avatar', '') or '',
                "style": "专业正式",  # 默认值
                "file_format": "PDF",  # 默认值
                "notify": True,
                "auto_save": True
            }
            
            # 2. 如果CSV中有额外字段，合并（CSV作为补充）
            file_path = "data/profiles.csv"
            if os.path.exists(file_path):
                with open(file_path, "r", encoding="utf-8-sig") as f:
                    reader = csv.DictReader(f)
                    for row in reader:
                        if row.get('username') == username:
                            # 合并CSV中的字段（如果数据库中没有）
                            if not profile_data.get('style') and row.get('style'):
                                profile_data['style'] = row.get('style', '专业正式')
                            if not profile_data.get('file_format') and row.get('file_format'):
                                profile_data['file_format'] = row.get('file_format', 'PDF')
                            if row.get('notify'):
                                profile_data['notify'] = row.get('notify') == 'True'
                            if row.get('auto_save'):
                                profile_data['auto_save'] = row.get('auto_save') == 'True'
                            break
            
            return {"success": True, "data": profile_data}
    except Exception as e:
        print(f"⚠️ [get_profile] 从数据库加载失败: {e}")
        # 继续尝试从CSV加载
    
    # 3. 如果数据库没有，从CSV加载
    file_path = "data/profiles.csv"
    if os.path.exists(file_path):
        with open(file_path, "r", encoding="utf-8-sig") as f:
            reader = csv.DictReader(f)
            for row in reader:
                if row.get('username') == username:
                    # 转换布尔值 (CSV里存的是字符串)
                    row['notify'] = row.get('notify') == 'True'
                    row['auto_save'] = row.get('auto_save') == 'True'
                    return {"success": True, "data": row}
    
    # 4. 都没找到，返回默认值
    return {"success": True, "data": {"username": username, "email": "", "phone": "", "city": "", "avatar": "", "style": "专业正式", "file_format": "PDF", "notify": True, "auto_save": True}}

# --- 3. 更新用户资料接口 ---
@app.post("/api/user/profile")
def update_profile(profile: UserProfile):
    file_path = "data/profiles.csv"
    os.makedirs("data", exist_ok=True)
    
    # 读取所有现存资料
    profiles = []
    if os.path.exists(file_path):
        with open(file_path, "r", encoding="utf-8-sig") as f:
            reader = csv.DictReader(f)
            profiles = list(reader)
    
    # 查找并更新，或者新增
    updated = False
    for row in profiles:
        if row.get('username') == profile.username:
            row.update(profile.dict()) # 更新字段 (这里会自动包含 avatar)
            # 把布尔值转回字符串存CSV
            row['notify'] = str(profile.notify)
            row['auto_save'] = str(profile.auto_save)
            updated = True
            break
    
    if not updated:
        # 新增一条
        new_row = profile.dict()
        new_row['notify'] = str(profile.notify)
        new_row['auto_save'] = str(profile.auto_save)
        profiles.append(new_row)
    
    # 写回 CSV 文件
    try:
        with open(file_path, "w", encoding="utf-8-sig", newline="") as f:
            fieldnames = ["username", "avatar", "email", "phone", "city", "style", "file_format", "notify", "auto_save"]
            writer = csv.DictWriter(f, fieldnames=fieldnames)
            writer.writeheader()
            writer.writerows(profiles)
        print(f"✅ [update_profile] CSV 文件更新成功")
    except Exception as e:
        print(f"⚠️ [update_profile] CSV 文件更新失败: {e}")
    
    # 2. 更新数据库（优先保存到数据库，确保持久化）
    try:
        # 使用 update_user_multiple_fields 更新数据库
        db_fields = {}
        # 即使字段为空，也保存（允许清空字段）
        db_fields['email'] = profile.email or ''
        db_fields['phone'] = profile.phone or ''
        db_fields['city'] = profile.city or ''
        if profile.avatar:
            db_fields['avatar'] = profile.avatar
        
        # 强制更新数据库，确保数据持久化
        success = update_user_multiple_fields(username, db_fields)
        if success:
            print(f"✅ [update_profile] 数据库更新成功: {db_fields}")
        else:
            print(f"⚠️ [update_profile] 数据库更新失败（可能字段不存在）")
            # 如果批量更新失败，尝试逐个字段更新
            for field, value in db_fields.items():
                try:
                    update_user_field(username, field, value)
                    print(f"✅ [update_profile] 字段 {field} 单独更新成功")
                except Exception as e:
                    print(f"⚠️ [update_profile] 字段 {field} 更新失败: {e}")
    except Exception as e:
        print(f"⚠️ [update_profile] 数据库更新异常: {e}")
        print(f"⚠️ [update_profile] 错误堆栈: {traceback.format_exc()}")
        # 数据库更新失败不影响 CSV 保存，继续返回成功
    
    return {"success": True, "message": "资料已保存"}
# ==========================================
#  核心功能 F: 简历医生 (诊断 + 生成)
# ==========================================
@app.post("/api/resume/analyze")
async def analyze_resume(file: UploadFile = File(...)):
    time.sleep(1.5)
    print(f"收到简历诊断请求: {file.filename}")
    
    # 核心：确保 score_rationale 存在
    return {
        "score": 82,
        "score_rationale": "✅ 基础分70。因项目使用了STAR法则+5分，技术栈匹配+10分；❌ 但缺少GitHub链接-3分。",
        "summary": "简历结构清晰，技术栈覆盖全面，但‘量化成果’有待提升。",
        "strengths": ["教育背景优秀", "两段相关实习", "技术栈命中率高"],
        "weaknesses": ["缺乏具体性能数据", "自我评价泛泛", "无开源贡献"],
        "suggestions": ["补充性能对比数据", "增加熟练度描述", "添加 GitHub 链接"]
    }

# --- 4. 修改密码接口 ---
class ChangePwdRequest(BaseModel):
    username: str
    old_password: str
    new_password: str

@app.post("/api/user/change_password")
def change_password(req: ChangePwdRequest):
    # 1. 验证旧密码（使用数据库）
    user = get_user_by_username(req.username)
    if not user:
        return {"success": False, "message": "用户不存在"}
    
    if user.get('password', '').strip() != req.old_password:
        return {"success": False, "message": "旧密码不正确"}
    
    # 2. 更新数据库密码
    success = update_user_field(req.username, "password", req.new_password)
    if success:
        return {"success": True, "message": "密码修改成功"}
    else:
        return {"success": False, "message": "密码更新失败"}

# --- 5. 上传头像接口 ---
@app.post("/api/user/avatar")
async def upload_avatar(
    file: UploadFile = File(...),  # 修改：使用 file 作为参数名，与前端 FormData 字段名匹配
    username: str = Form(...)
):
    """
    上传用户头像接口
    
    接收图片文件，保存到服务器，返回可访问的 URL
    同时更新数据库和 CSV 文件中的 avatar 字段
    
    注意：前端 FormData 字段名应为 'file'，不是 'avatar'
    """
    import traceback
    import uuid
    from datetime import datetime
    
    print(f"✅ [upload_avatar] 收到头像上传请求，用户: {username}, 文件名: {file.filename}")
    
    # 1. 验证用户是否存在
    user = get_user_by_username(username)
    if not user:
        raise HTTPException(status_code=404, detail="用户不存在")
    
    # 2. 验证文件类型
    if not file.filename:
        raise HTTPException(status_code=400, detail="文件名不能为空")
    
    file_ext = os.path.splitext(file.filename)[1].lower()
    allowed_exts = ['.jpg', '.jpeg', '.png', '.gif', '.webp']
    if file_ext not in allowed_exts:
        raise HTTPException(status_code=400, detail=f"不支持的文件格式，仅支持: {', '.join(allowed_exts)}")
    
    # 3. 验证文件大小（限制 10MB，已解除更严格的限制）
    file_content = await file.read()
    file_size_mb = len(file_content) / (1024 * 1024)
    print(f"📊 [upload_avatar] 文件大小: {file_size_mb:.2f} MB")
    
    if len(file_content) > 10 * 1024 * 1024:
        raise HTTPException(status_code=413, detail=f"文件大小不能超过 10MB，当前文件大小: {file_size_mb:.2f} MB")
    
    # 4. 生成唯一文件名（避免冲突）
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    unique_id = str(uuid.uuid4())[:8]
    safe_filename = f"{username}_{timestamp}_{unique_id}{file_ext}"
    file_path = os.path.join("static", "avatars", safe_filename)
    
    # 确保目录存在
    os.makedirs(os.path.dirname(file_path), exist_ok=True)
    
    # 5. 保存文件
    try:
        with open(file_path, "wb") as buffer:
            buffer.write(file_content)
        print(f"✅ [upload_avatar] 文件保存成功: {file_path}")
    except Exception as e:
        print(f"❌ [upload_avatar] 文件保存失败: {e}")
        raise HTTPException(status_code=500, detail=f"文件保存失败: {str(e)}")
    
    # 6. 生成可访问的 URL（使用相对路径，前端会拼接 API_BASE）
    # 注意：使用相对路径，前端会根据 API_BASE 自动拼接完整 URL
    avatar_url = f"/static/avatars/{safe_filename}"
    
    # 7. 更新数据库中的 avatar 字段（使用完整 URL 存储）
    try:
        # 生成完整 URL 用于数据库存储
        base_url = os.getenv("BASE_URL", "https://ai-career-helper-backend-u1s0.onrender.com")
        full_avatar_url = f"{base_url}{avatar_url}"
        
        success = update_user_field(username, "avatar", full_avatar_url)
        if success:
            print(f"✅ [upload_avatar] 数据库 avatar 字段更新成功: {full_avatar_url}")
        else:
            print(f"⚠️ [upload_avatar] 数据库 avatar 字段更新失败（可能字段不存在）")
    except Exception as e:
        print(f"⚠️ [upload_avatar] 数据库更新异常: {e}")
        print(f"⚠️ [upload_avatar] 错误堆栈: {traceback.format_exc()}")
        # 数据库更新失败不影响文件上传，继续返回成功
    
    # 8. 更新 CSV 文件中的 avatar 字段（保持兼容）
    try:
        csv_path = "data/profiles.csv"
        os.makedirs(os.path.dirname(csv_path), exist_ok=True)  # 确保目录存在
        
        # 生成完整 URL 用于 CSV 存储
        base_url = os.getenv("BASE_URL", "https://ai-career-helper-backend-u1s0.onrender.com")
        full_avatar_url = f"{base_url}{avatar_url}"
        
        profiles = []
        if os.path.exists(csv_path):
            with open(csv_path, "r", encoding="utf-8-sig") as f:
                reader = csv.DictReader(f)
                profiles = list(reader)
        
        updated = False
        for row in profiles:
            if row.get('username') == username:
                row['avatar'] = full_avatar_url
                updated = True
                break
        
        if not updated:
            # 如果用户资料不存在，创建新记录
            profiles.append({
                'username': username,
                'avatar': full_avatar_url,
                'email': '',
                'phone': '',
                'city': '',
                'style': '专业正式',
                'file_format': 'PDF',
                'notify': 'True',
                'auto_save': 'True'
            })
        
        with open(csv_path, "w", encoding="utf-8-sig", newline="") as f:
            fieldnames = ["username", "avatar", "email", "phone", "city", "style", "file_format", "notify", "auto_save"]
            writer = csv.DictWriter(f, fieldnames=fieldnames)
            writer.writeheader()
            writer.writerows(profiles)
        print(f"✅ [upload_avatar] CSV 文件 avatar 字段更新成功: {full_avatar_url}")
    except Exception as e:
        print(f"⚠️ [upload_avatar] CSV 文件更新异常: {e}")
        print(f"⚠️ [upload_avatar] CSV 错误堆栈: {traceback.format_exc()}")
        # CSV 更新失败不影响主流程
    
    # 9. 返回结果（返回完整 URL，前端可以直接使用）
    base_url = os.getenv("BASE_URL", "https://ai-career-helper-backend-u1s0.onrender.com")
    full_avatar_url = f"{base_url}{avatar_url}"
    
    return {
        "success": True,
        "avatar_url": full_avatar_url,
        "avatar": full_avatar_url,  # 兼容字段名
        "url": full_avatar_url,  # 兼容字段名
        "message": "头像上传成功"
    }
@app.post("/api/resume/generate")
def generate_resume(req: GenerateResumeRequest):
    time.sleep(1.5)
    direction = req.focus_direction
    print(f"收到生成请求，方向: {direction}")
    
    # 动态生成内容
    if "算法" in direction:
        tech = "PyTorch, Transformer, CUDA"
        role = "算法工程师"
    elif "前端" in direction:
        tech = "Vue3, TypeScript, Vite, Element Plus"
        role = "前端开发工程师"
    else:
        tech = "FastAPI, Vue3, Docker, Redis"
        role = "全栈开发工程师"

    content = f"""
# 你的姓名 (意向岗位：{role})
电话：138-xxxx-xxxx | 邮箱：email@example.com

## 💡 AI 优化摘要
> **优化重点**：根据 **{direction}** 方向重构了技能清单，并引入 **STAR 法则** 优化了项目描述。

## 🎓 教育背景
**北京邮电大学** | 人工智能学院 | 本科 | 2024-2028
* **主修课程**：数据结构 (95)、机器学习 (92)
* **核心优势**：专业排名前 10%

## 💻 项目经历 (精修版)
**AI 简历全科医生平台** | 全栈负责人 | {tech}
* **背景 (S)**：针对大学生求职痛点，开发智能辅助系统。
* **任务 (T)**：负责从 0 到 1 搭建前后端分离架构。
* **行动 (A)**：
    * **架构设计**：基于 **FastAPI** 重构接口，修复了“404 Not Found"的关键 Bug。
    * **体验优化**：前端采用 **Vue3** 实现“双屏联动”，效率提升 **50%"。
* **结果 (R)**：项目上线首周获得 200+ 次调用。

## 🛠 技能清单
* **核心技术**：{tech}
* **工具**：Git, Linux

## 📜 自我评价
* 具备极强的 Debug 能力，善于在压力下快速定位并解决问题。
"""
    return {"success": True, "content": content.strip()}


# ===================== 新增：简历上传相关接口 =====================
class ResumeUploadRequest(BaseModel):
    username: str
    task_id: str
    filename: str
    report: str
    score: float | int = 0
    date: str | None = None


@app.post('/api/resume/upload')
def upload_resume(item: ResumeUploadRequest):
    """接收前端上传的简历报告，持久化到 data/uploaded_resumes.csv 并更新数据库的 uploadedResumeNum 字段"""
    os.makedirs('data', exist_ok=True)
    uploaded_file = 'data/uploaded_resumes.csv'
    users_file = 'data/users.csv'

    # 填充默认日期
    if not item.date:
        item.date = datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')

    # 1. 写入 uploaded_resumes.csv
    fieldnames = ['task_id', 'username', 'filename', 'report', 'score', 'date']
    exists = os.path.exists(uploaded_file)
    try:
        with open(uploaded_file, 'a', encoding='utf-8-sig', newline='') as f:
            writer = csv.DictWriter(f, fieldnames=fieldnames)
            if not exists:
                writer.writeheader()
            writer.writerow({
                'task_id': item.task_id,
                'username': item.username,
                'filename': item.filename,
                'report': item.report,
                'score': item.score,
                'date': item.date,
            })
    except Exception as e:
        return {'success': False, 'message': f'写入上传记录失败: {e}'}

    # 2. 更新数据库的 uploadedResumeNum 字段
    try:
        success = increment_user_field(item.username, "uploadedResumeNum", 1)
        if not success:
            print(f"⚠️ [DEBUG] 用户 {item.username} 的 uploadedResumeNum 更新失败")
    except Exception as e:
        print(f"⚠️ [DEBUG] 更新 uploadedResumeNum 失败: {e}")
        # 不阻碍上传，但记录提示
        return {'success': True, 'message': '上传成功，但用户统计更新失败'}

    return {'success': True, 'message': '上传成功'}


@app.get('/api/resume/getUploadedList')
def get_uploaded_list():
    """返回所有已上传的简历上传记录（若无则生成 3 条模拟数据）"""
    os.makedirs('data', exist_ok=True)
    uploaded_file = 'data/uploaded_resumes.csv'

    # 如果文件不存在，生成三条默认模拟数据
    if not os.path.exists(uploaded_file):
        mock = [
            {'task_id': 'T-MOCK-01', 'username': 'alice', 'filename': 'alice_resume.pdf', 'report': '# 模拟报告\n- 分数：88\n- 建议：突出项目', 'score': 88, 'date': datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')},
            {'task_id': 'T-MOCK-02', 'username': 'bob', 'filename': 'bob_resume.pdf', 'report': '# 模拟报告\n- 分数：76\n- 建议：补充实习', 'score': 76, 'date': datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')},
            {'task_id': 'T-MOCK-03', 'username': 'carol', 'filename': 'carol_resume.pdf', 'report': '# 模拟报告\n- 分数：92\n- 建议：保持精炼', 'score': 92, 'date': datetime.datetime.now().strftime('%Y-%m-%d %H:%M:%S')},
        ]
        with open(uploaded_file, 'w', encoding='utf-8-sig', newline='') as f:
            writer = csv.DictWriter(f, fieldnames=['task_id','username','filename','report','score','date'])
            writer.writeheader()
            writer.writerows(mock)

    # 读取并返回
    records = []
    with open(uploaded_file, 'r', encoding='utf-8-sig') as f:
        reader = csv.DictReader(f)
        for row in reader:
            # convert numeric
            row['score'] = float(row.get('score') or 0)
            records.append(row)
    # 返回按时间倒序（最新在前）
    records.reverse()
    return {'success': True, 'data': records}


@app.post('/api/resume/delete')
def delete_upload(username: str, task_id: str):
    """删除上传记录并同步数据库的统计字段"""
    uploaded_file = 'data/uploaded_resumes.csv'

    if not os.path.exists(uploaded_file):
        return {'success': False, 'message': '没有上传记录文件'}

    rows = []
    removed = False
    with open(uploaded_file, 'r', encoding='utf-8-sig') as f:
        reader = csv.DictReader(f)
        for row in reader:
            if row.get('username') == username and row.get('task_id') == task_id:
                removed = True
                continue
            rows.append(row)

    if removed:
        with open(uploaded_file, 'w', encoding='utf-8-sig', newline='') as f:
            writer = csv.DictWriter(f, fieldnames=['task_id','username','filename','report','score','date'])
            writer.writeheader()
            writer.writerows(rows)

        # 同步数据库 uploadedResumeNum 减一
        try:
            decrement_user_field(username, "uploadedResumeNum", 1)
        except Exception as e:
            print(f"⚠️ [DEBUG] 更新 uploadedResumeNum 失败: {e}")

        return {'success': True, 'message': '删除上传记录成功'}

    return {'success': False, 'message': '未找到对应上传记录'}


@app.post('/api/user/addTask')
def add_user_task(username: str):
    """为用户的 createTaskNum +1（用于统计用户提交到 Admin 的次数）"""
    try:
        success = increment_user_field(username, "createTaskNum", 1)
        if success:
            return {'success': True, 'message': '用户任务数已更新'}
        else:
            return {'success': False, 'message': '未找到用户或更新失败'}
    except Exception as e:
        print(f"❌ [DEBUG] 更新 createTaskNum 失败: {e}")
        return {'success': False, 'message': f'更新失败: {e}'}
# ==========================================
# 🎮 虚拟职业体验模块 (Career Simulation)
# ==========================================

# 1. 模拟剧本数据 (Mock Data)
SIMULATION_SCRIPTS = {
    "product_manager": {
        "title": "产品经理的一天",
        "desc": "体验从需求评审到上线发布的生死时速。",
        "scenes": [
            {
                "id": 1,
                "text": "早上9:30，你刚到公司，开发组长气冲冲地跑过来说：'昨天定的需求技术实现不了，必须砍掉！' 同时，运营那边催着要上线。你会怎么做？",
                "options": [
                    {"label": "坚持原需求，让开发想办法", "score_change": -10, "feedback": "开发组长拍了桌子，项目延期风险增加。"},
                    {"label": "立刻砍掉功能，保上线", "score_change": 5, "feedback": "运营很不满，但至少能按时上线。"},
                    {"label": "拉会协调，寻找替代方案", "score_change": 10, "feedback": "虽然花了一小时开会，但大家达成了共识，干得漂亮！"}
                ]
            },
            {
                "id": 2,
                "text": "下午3:00，老板突然在群里发了一张竞品的截图，说：'这个功能很酷，我们要不要也加一个？' 此时距离封版只剩2小时。",
                "options": [
                    {"label": "老板说加就加！", "score_change": -20, "feedback": "开发全员炸锅，今晚通宵已成定局，士气低落。"},
                    {"label": "私聊老板，说明风险，建议下个版本加", "score_change": 10, "feedback": "老板觉得你考虑周全，同意了你的建议。"},
                    {"label": "装作没看见", "score_change": -5, "feedback": "老板在群里@了你，场面一度十分尴尬。"}
                ]
            }
        ]
    },
    "programmer": {
        "title": "全栈工程师的一天",
        "desc": "体验代码、Bug与产品经理之间的爱恨情仇。",
        "scenes": [
            {
                "id": 1,
                "text": "上午10:00，你正在写核心代码，产品经理突然跑过来说：'这个按钮的颜色能不能换成五彩斑斓的黑？' 你被打断了思路。",
                "options": [
                    {"label": "直接怼回去：'你行你上！'", "score_change": -10, "feedback": "产品经理哭着去找老板了，你被HR约谈。"},
                    {"label": "耐心解释技术实现难度", "score_change": 10, "feedback": "产品经理被你的专业术语绕晕了，放弃了修改。"},
                    {"label": "默默记下，先写完手头代码", "score_change": 5, "feedback": "稳妥的做法，但需求还是得改。"}
                ]
            },
            {
                "id": 2,
                "text": "下午5:50，准备下班去约会。测试突然提了一个 '严重' 级别的Bug，说是偶发性的，复现不出来。",
                "options": [
                    {"label": "不管了，先下班", "score_change": -15, "feedback": "线上炸了，你在约会途中被叫回公司修通宵。"},
                    {"label": "留下来排查，推迟约会", "score_change": 10, "feedback": "查出了是缓存问题，半小时搞定，不仅没迟到还收获了测试的崇拜。"},
                    {"label": "告诉测试：'我本地是好的'", "score_change": -5, "feedback": "经典的程序员语录，但问题依然存在。"}
                ]
            }
        ]
    }
}

class SimulationRequest(BaseModel):
    role_id: str

# 2. 获取剧本接口
@app.post("/api/simulation/start")
def start_simulation(req: SimulationRequest):
    role = req.role_id
    if role not in SIMULATION_SCRIPTS:
        return {"success": False, "message": "剧本不存在"}
    
    script = SIMULATION_SCRIPTS[role]
    return {
        "success": True, 
        "data": {
            "title": script["title"],
            "scenes": script["scenes"] # 一次性把简单剧本都给前端，前端自己控制进度
        }
    }

# ==========================================
#  新增功能 G: 虚拟职业体验 & 生涯分析整合
# ==========================================
@app.post("/api/virtual-career/questions")
def virtual_career_questions(req: VirtualCareerQuestionsRequest):
    """
    根据职业名称动态生成 15 道匹配度选择题（每题 4 个选项）
    """
    system_prompt = (
        "你是一名职业规划评估题目设计专家。"
        "请针对指定职业设计 15 道用于评估匹配度的单选题，每题 4 个选项。"
        "题目要尽量贴近真实工作场景，覆盖能力要求、工作方式偏好、压力/节奏、沟通协作等维度。\n"
        "必须严格按照以下 JSON 结构返回：\n"
        "{\n"
        '  \"career\": \"职业名称\",\n'
        '  \"questions\": [\n'
        "    {\"id\": \"q1\", \"title\": \"题目 1 文本\", \"options\": [\"选项A\", \"选项B\", \"选项C\", \"选项D\"]},\n"
        "    ... 共 15 道题 ...\n"
        "  ]\n"
        "}"
    )

    user_prompt = (
        "目标职业名称：\n"
        f"{req.career}\n\n"
        "如果这是一个非常冷门或未见过的职业，请先用 1-2 句话理解/假设这个职业的核心工作内容，"
        "然后基于你的理解设计题目。"
    )

    data = _deepseek_json(system_prompt, user_prompt)
    questions = data.get("questions") or []
    if not isinstance(questions, list) or len(questions) == 0:
        raise HTTPException(status_code=500, detail="AI 生成题目失败，请稍后重试")

    # 简单保证 id 存在
    for idx, q in enumerate(questions, start=1):
        q.setdefault("id", f"q{idx}")

    return {
        "career": data.get("career", req.career),
        "questions": questions[:15],
    }

@app.post("/api/generate-job-test")
def generate_job_test(req: GenerateJobTestRequest):
    """
    根据用户输入的职业名称：
    1. 调用 DeepSeek 生成该职业的虚拟体验脚本（包含职业定义、典型场景、3~5 个互动选择及结果）
    2. 再调用 DeepSeek 生成 15 道职业相关测试题（单选题，每题 4 个选项）

    返回格式（与原有接口 /api/virtual-career/questions 完全一致）：
    {
      "jobScript": "AI生成的职业体验脚本",
      "questions": [
        {
          "id": "q1",
          "title": "题干",
          "options": ["选项A", "选项B", "选项C", "选项D"],
          "correctAnswer": "A",
          "score": 1
        },
        ... 共 15 道题 ...
      ]
    }

    异常格式：
    - 职业名为空：{code: 400, msg: "请输入职业名"}
    - AI 生成失败：{code: 500, msg: "AI生成失败，请稍后重试"}
    
    注意：此接口仅支持 POST 请求方法，已配置 CORS 支持 OPTIONS 预检请求
    """
    import traceback
    print(f"✅ [generate-job-test] 收到 POST 请求: jobName={req.jobName}")
    print(f"✅ [generate-job-test] 请求体类型: {type(req)}, 请求体内容: {req.model_dump()}")
    
    job_name = (req.jobName or "").strip()
    if not job_name:
        print(f"❌ [generate-job-test] 职业名为空")
        return JSONResponse(status_code=400, content={"code": 400, "msg": "请输入职业名"})

    try:
        print(f"🔄 [generate-job-test] 开始生成职业体验脚本: {job_name}")
        # 第一步：生成职业体验脚本（文本即可，可包含 Markdown）
        script_system_prompt = (
            "你是一名职业体验设计师，负责为用户设计沉浸式「虚拟职业体验」脚本。\n"
            "目标：针对指定职业，生成一段完整的体验脚本，帮助用户在几分钟内沉浸式感受该职业的真实工作场景。\n"
            "必须包含以下内容（使用清晰的小标题或分段）：\n"
            "1. 职业定义与核心职责概述（1~2 段）\n"
            "2. 典型的一天/一周工作场景（2~3 段，可以具体到时间点和任务）\n"
            "3. 设计 3~5 个关键抉择节点，每个节点：\n"
            "   - 先用 2~3 句话描述当前情境\n"
            "   - 给出 3 个左右可选操作（用 A/B/C 编号）\n"
            "   - 对每个选项给出简短的结果反馈（包括积极或消极影响）\n"
            "4. 最后的总结与建议（根据用户在体验中的倾向，给出 3~5 条建议）\n"
            "要求：\n"
            "- 使用通俗易懂的中文，语气亲切、有画面感\n"
            "- 可以使用 Markdown 标题/列表增强可读性，但不要输出任何 JSON 结构\n"
        )
        script_user_prompt = (
            f"目标职业名称：{job_name}\n\n"
            "请基于你对该职业的理解，按照上述结构输出完整的职业体验脚本。"
        )
        script_text = _deepseek_markdown(script_system_prompt, script_user_prompt)
        print(f"✅ [generate-job-test] 职业体验脚本生成完成，长度: {len(script_text)} 字符")

        print(f"🔄 [generate-job-test] 开始生成 15 道测试题: {job_name}")
        # 第二步：生成 15 道职业测试题（JSON）
        # 返回格式必须与原有接口 /api/virtual-career/questions 完全一致
        questions_system_prompt = (
            "你是一名职业规划评估题目设计专家。\n"
            "请针对指定职业设计 15 道用于评估匹配度的单选题，每题 4 个选项。\n"
            "题目要尽量贴近真实工作场景，覆盖能力要求、工作方式偏好、压力/节奏、沟通协作等维度。\n"
            "必须严格按照以下 JSON 结构返回：\n"
            "{\n"
            "  \"questions\": [\n"
            "    {\n"
            "      \"title\": \"题目 1 文本\",\n"
            "      \"options\": [\"选项A\", \"选项B\", \"选项C\", \"选项D\"],\n"
            "      \"correctAnswer\": \"A\",\n"
            "      \"score\": 1\n"
            "    },\n"
            "    ... 共 15 道题 ...\n"
            "  ]\n"
            "}\n"
            "注意：\n"
            "- correctAnswer 必须是单个选项字母（如 \"A\"），表示正确答案\n"
            "- score 为每题分值，统一为 1\n"
            "- 确保每道题都有 4 个选项"
        )
        questions_user_prompt = (
            "目标职业名称：\n"
            f"{job_name}\n\n"
            "如果这是一个非常冷门或未见过的职业，请先用 1-2 句话理解/假设这个职业的核心工作内容，"
            "然后基于你的理解设计题目。"
        )

        questions_data = _deepseek_json(questions_system_prompt, questions_user_prompt)
        raw_questions = questions_data.get("questions") or []
        print(f"✅ [generate-job-test] AI 返回原始题目数量: {len(raw_questions)}")

        # 基本校验
        if not isinstance(raw_questions, list) or len(raw_questions) == 0:
            raise ValueError("AI 未生成有效题目")

        # 归一化为与原有接口完全一致的格式
        normalized_questions = []
        for idx, q in enumerate(raw_questions[:15], start=1):
            if not isinstance(q, dict):
                continue
            
            # 提取字段，兼容多种可能的字段名
            question_text = q.get("title") or q.get("question") or q.get("stem") or f"第 {idx} 题"
            options = q.get("options") or []
            correct_answer = q.get("correctAnswer") or q.get("answer") or q.get("correct") or ""
            score = q.get("score")
            
            # 确保选项为字符串列表，且至少有 4 个选项
            options = [str(o) for o in options]
            while len(options) < 4:
                options.append(f"选项{chr(68 + len(options))}")  # 补充到 4 个选项
            
            # 确保 correctAnswer 是单个字母（如 "A"）
            if correct_answer:
                # 如果答案是 "A"、"B" 等，直接使用；如果是 "选项A"，提取字母
                if len(correct_answer) == 1 and correct_answer.isalpha():
                    correct_answer = correct_answer.upper()
                elif "选项" in correct_answer or correct_answer.startswith("选项"):
                    # 尝试从 "选项A" 中提取 "A"
                    for char in correct_answer:
                        if char.isalpha():
                            correct_answer = char.upper()
                            break
                else:
                    # 默认取第一个字符
                    correct_answer = str(correct_answer)[0].upper() if correct_answer else "A"
            else:
                correct_answer = "A"  # 默认答案
            
            # score 默认为 1
            if score is None:
                score = 1
            else:
                try:
                    score = int(score)
                except (ValueError, TypeError):
                    score = 1

            # 构建与原有接口完全一致的题目结构
            normalized_questions.append(
                {
                    "id": f"q{idx}",  # 保持字符串格式以兼容前端（前端使用 q.id.toUpperCase()）
                    "title": question_text,
                    "options": options[:4],  # 确保只有 4 个选项
                    "correctAnswer": correct_answer,
                    "score": score,
                }
            )

        if not normalized_questions:
            raise ValueError("AI 生成的题目结构异常")

        # 返回格式与原有接口对齐：使用 jobScript 字段名（与原有 script 字段对应）
        result = {
            "jobScript": script_text,
            "questions": normalized_questions,
        }
        print(f"✅ [generate-job-test] 成功生成 {len(normalized_questions)} 道题目，准备返回结果")
        return result

    except Exception as e:
        # 统一转为前端友好的错误结构
        print(f"❌ [generate-job-test] 生成失败: {e}")
        print(f"❌ [generate-job-test] 错误堆栈: {traceback.format_exc()}")
        return JSONResponse(status_code=500, content={"code": 500, "msg": "AI生成失败，请稍后重试"})


@app.post("/api/analyze_resume")
async def analyze_resume(
    resume_file: UploadFile = File(...),
    resume_text: Optional[str] = Form(None),
):
    """
    简历诊断与优化接口
    
    支持两种输入方式：
    1. 文件上传：resume_file（支持 PDF/DOCX/TXT）
    2. 文本输入：resume_text（直接传入简历文本）
    
    返回格式：
    {
      "success": true,
      "diagnosis_report": {
        "score": 85,
        "summary": "AI生成的综合评价",
        "score_details": ["评分依据1", "评分依据2"],
        "highlights": ["亮点1", "亮点2"],
        "weaknesses": ["不足1", "不足2"]
      },
      "optimized_resume": "AI生成的Markdown格式优化简历",
      "fallback": false
    }
    """
    import traceback
    
    print(f"✅ [analyze_resume] 收到简历分析请求")
    resume_file_name = resume_file.filename if resume_file and hasattr(resume_file, 'filename') else None
    print(f"✅ [analyze_resume] 参数: resume_file={resume_file_name}, resume_text={'已提供' if resume_text else None}")
    
    # 1. 提取简历文本内容
    resume_content = ""
    try:
        if resume_file:
            print(f"🔄 [analyze_resume] 开始解析文件: {resume_file.filename}")
            resume_content = extract_text_from_file(resume_file)
        elif resume_text:
            print(f"🔄 [analyze_resume] 使用文本输入，长度: {len(resume_text)} 字符")
            resume_content = resume_text.strip()
        else:
            return JSONResponse(
                status_code=400,
                content={"success": False, "error": "请提供简历文件（resume_file）或简历文本（resume_text）"}
            )
        
        if not resume_content:
            return JSONResponse(
                status_code=400,
                content={"success": False, "error": "简历内容为空，请检查文件或文本内容"}
            )
        
        print(f"✅ [analyze_resume] 简历内容提取成功，长度: {len(resume_content)} 字符")
    
    except HTTPException:
        raise
    except Exception as e:
        print(f"❌ [analyze_resume] 文件解析异常: {e}")
        print(f"❌ [analyze_resume] 错误堆栈: {traceback.format_exc()}")
        return JSONResponse(
            status_code=500,
            content={"success": False, "error": f"文件解析失败: {str(e)}"}
        )
    
    # 2. 调用 DeepSeek 生成诊断报告和优化简历
    fallback_used = False
    diagnosis_report = None
    optimized_resume = None
    
    try:
        print(f"🔄 [analyze_resume] 开始调用 DeepSeek API 生成诊断报告")
        
        # 2.1 生成诊断报告
        diagnosis_system_prompt = (
            "你是资深简历优化专家，分析以下简历内容，严格按以下JSON结构输出诊断报告，不要任何多余话术：\n"
            "{\n"
            '  "score": 数字（0-100）,\n'
            '  "summary": "综合评价一句话",\n'
            '  "score_details": ["评分依据1", "评分依据2"],\n'
            '  "highlights": ["亮点1", "亮点2"],\n'
            '  "weaknesses": ["不足1", "不足2"]\n'
            "}"
        )
        diagnosis_user_prompt = f"简历内容：\n{resume_content}"
        
        diagnosis_data = _deepseek_json(diagnosis_system_prompt, diagnosis_user_prompt)
        
        # 归一化诊断报告结构
        diagnosis_report = {
            "score": int(diagnosis_data.get("score", 0)) if isinstance(diagnosis_data.get("score"), (int, float)) else 0,
            "summary": diagnosis_data.get("summary", "AI 暂未生成综合评价"),
            "score_details": diagnosis_data.get("score_details", []) if isinstance(diagnosis_data.get("score_details"), list) else [],
            "highlights": diagnosis_data.get("highlights", []) if isinstance(diagnosis_data.get("highlights"), list) else [],
            "weaknesses": diagnosis_data.get("weaknesses", []) if isinstance(diagnosis_data.get("weaknesses"), list) else [],
        }
        
        print(f"✅ [analyze_resume] 诊断报告生成成功，评分: {diagnosis_report['score']}")
        
        # 2.2 生成优化简历
        print(f"🔄 [analyze_resume] 开始调用 DeepSeek API 生成优化简历")
        
        optimize_system_prompt = (
            "基于以下简历内容，优化为更专业的版本，严格按以下Markdown结构输出，不要任何多余话术：\n"
            "# 你的姓名 (意向岗位: 全栈开发工程师)\n"
            "电话: 138-xxxx-xxxx | 邮箱: email@example.com\n\n"
            "## 💡 AI优化摘要\n"
            "优化重点: ...\n\n"
            "## 🎓 教育背景\n"
            "北京邮电大学 | 人工智能学院 | 本科 | 2024-2028\n"
            "- 主修课程: ...\n"
            "- 核心优势: ...\n\n"
            "## 💻 项目经历 (精修版)\n"
            "### AI简历全科医生平台 | 全栈负责人 | FastAPI, Vue3, Docker, Redis\n"
            "- **背景(S)**: ...\n"
            "- **任务(T)**: ...\n"
            "- **行动(A)**: ...\n"
            "- **结果(R)**: ...\n\n"
            "## 🛠️ 技能清单\n"
            "- 核心技术: ...\n"
            "- 工具: ...\n\n"
            "## 📄 自我评价\n"
            "- ..."
        )
        optimize_user_prompt = f"简历内容：\n{resume_content}"
        
        optimized_resume = _deepseek_markdown(optimize_system_prompt, optimize_user_prompt)
        
        print(f"✅ [analyze_resume] 优化简历生成成功，长度: {len(optimized_resume)} 字符")
    
    except Exception as e:
        print(f"❌ [analyze_resume] DeepSeek API 调用失败: {e}")
        print(f"❌ [analyze_resume] 错误堆栈: {traceback.format_exc()}")
        
        # 降级逻辑：返回预设的诊断报告和优化简历
        fallback_used = True
        print(f"⚠️ [analyze_resume] 启用降级逻辑，返回预设内容")
        
        diagnosis_report = {
            "score": 82,
            "summary": "简历结构清晰，技术栈覆盖全面，但「量化成果」有待提升。",
            "score_details": [
                "✅ 基础分70。因项目使用了STAR法则+5分，技术栈匹配+10分；❌ 但缺少GitHub链接-3分。"
            ],
            "highlights": [
                "教育背景优秀",
                "两段相关实习",
                "技术栈命中率高"
            ],
            "weaknesses": [
                "缺乏具体性能数据",
                "自我评价泛泛",
                "无开源贡献"
            ]
        }
        
        optimized_resume = (
            "# 优化简历（降级模式）\n\n"
            "## 💡 AI优化摘要\n"
            "优化重点: 基于原始简历内容进行结构化优化，突出技术能力和项目成果。\n\n"
            "## 🎓 教育背景\n"
            "（请根据实际简历内容填写）\n\n"
            "## 💻 项目经历 (精修版)\n"
            "（请使用STAR法则重构项目描述）\n\n"
            "## 🛠️ 技能清单\n"
            "（请列出核心技术栈和工具）\n\n"
            "## 📄 自我评价\n"
            "（请补充具体的能力描述和职业目标）\n"
        )
    
    # 3. 返回结果
    return {
        "success": True,
        "diagnosis_report": diagnosis_report,
        "optimized_resume": optimized_resume,
        "fallback": fallback_used
    }


@app.post("/api/analyze-experiment")
def analyze_experiment(req: AnalyzeExperimentRequest):
    """
    接收 15 题答案字典或竞争力沙盘数据，调用 DeepSeek 生成 Markdown 分析报告
    
    支持两种场景：
    1. 虚拟职业体验：接收 15 题答案字典
    2. 竞争力沙盘：接收 6 维度原始输入 + 量化后的雷达图数据
    """
    target_career = req.career or "未指定（请根据答题推断最匹配的方向）"
    
    # 判断是否为竞争力沙盘分析请求
    is_competitiveness_sandbox = (
        target_career == "个人竞争力沙盘分析" or
        "竞争力沙盘" in target_career or
        ("雷达图量化数据" in str(req.answers) if isinstance(req.answers, dict) else False)
    )
    
    if is_competitiveness_sandbox:
        # 竞争力沙盘分析：基于 6 维度原始输入和量化数据生成报告
        system_prompt = (
            "你是一位资深职业竞争力评估专家和生涯规划顾问。\n"
            "用户提供了 6 个维度的竞争力数据（原始输入和量化后的 0-100 分数）。\n"
            "请基于这些数据，生成一份详细的个人竞争力分析报告（Markdown 格式），包含：\n"
            "1) **竞争力总览**：整体竞争力评级（优秀/良好/中等/待提升）\n"
            "2) **各维度深度分析**：\n"
            "   - GPA/学业成绩：分析学术能力水平及提升建议\n"
            "   - 项目实战经验：评估项目能力与岗位匹配度\n"
            "   - 名企实习经历：分析实习价值与职业准备度\n"
            "   - 竞赛获奖情况：评估竞赛含金量与竞争力\n"
            "   - 英语学术能力：分析英语水平与国际化潜力\n"
            "   - 领导力与协作：评估团队协作与领导潜力\n"
            "3) **优势亮点**：列出 3-5 个最突出的优势维度\n"
            "4) **短板与提升方向**：指出 3-5 个需要重点提升的维度，并给出具体建议\n"
            "5) **综合提升建议**：基于整体竞争力，提供 4-8 周的针对性提升计划\n"
            "6) **职业匹配建议**：根据竞争力数据，推荐 3-5 个匹配的职业方向\n"
            "要求：\n"
            "- 结合用户提供的原始输入（如\"GPA 3.8/4.0\"、\"字节1个月实习\"、\"省二\"等）进行具体分析\n"
            "- 结合量化后的雷达图分数，给出数据支撑的结论\n"
            "- 输出格式为 Markdown，使用标题、列表、加粗等格式增强可读性\n"
            "- 语言专业但易懂，避免过于学术化的表述"
        )
        
        user_prompt = (
            f"以下是用户的竞争力数据：\n\n"
            f"{json.dumps(req.answers, ensure_ascii=False, indent=2)}\n\n"
            f"请基于以上数据，生成一份详细的个人竞争力分析报告。"
        )
        
        try:
            markdown = _deepseek_markdown(system_prompt, user_prompt)
            
            # 提取量化后的分数（如果存在）
            quantized_scores = None
            if isinstance(req.answers, dict) and "雷达图量化数据(0-100)" in req.answers:
                quantized_scores = req.answers["雷达图量化数据(0-100)"]
            
            return {
                "success": True,
                "markdown": markdown,
                "analysis_report": markdown,  # 兼容前端期望的字段名
                "quantized_scores": quantized_scores,
                "fallback": False
            }
        except Exception as e:
            print(f"❌ [analyze_experiment] 竞争力沙盘分析失败: {e}")
            import traceback
            print(traceback.format_exc())
            # 降级：返回基础报告
            fallback_report = (
                "## 📊 竞争力总览\n\n"
                "基于您提供的 6 个维度数据，系统已进行初步分析。\n\n"
                "## 📈 各维度分数\n\n"
                f"- GPA 学术成绩：{req.answers.get('雷达图量化数据(0-100)', {}).get('gpa', 0) if isinstance(req.answers, dict) else 0}/100\n"
                f"- 项目实战经验：{req.answers.get('雷达图量化数据(0-100)', {}).get('project', 0) if isinstance(req.answers, dict) else 0}/100\n"
                f"- 名企实习经历：{req.answers.get('雷达图量化数据(0-100)', {}).get('intern', 0) if isinstance(req.answers, dict) else 0}/100\n"
                f"- 竞赛获奖情况：{req.answers.get('雷达图量化数据(0-100)', {}).get('competition', 0) if isinstance(req.answers, dict) else 0}/100\n"
                f"- 英语学术能力：{req.answers.get('雷达图量化数据(0-100)', {}).get('english', 0) if isinstance(req.answers, dict) else 0}/100\n"
                f"- 领导力与协作：{req.answers.get('雷达图量化数据(0-100)', {}).get('leader', 0) if isinstance(req.answers, dict) else 0}/100\n\n"
                "## 💡 建议\n\n"
                "建议重点关注分数较低的维度，制定针对性的提升计划。"
            )
            return {
                "success": True,
                "markdown": fallback_report,
                "analysis_report": fallback_report,
                "quantized_scores": quantized_scores,
                "fallback": True
            }
    else:
        # 虚拟职业体验：15 题答案分析
        system_prompt = (
            "你是一位资深生涯规划师与组织心理学顾问。"
            "用户针对某一职业完成了 15 道匹配度选择题（每题 4 个选项）。"
            "请为该用户生成一份围绕\"目标职业匹配度\"的 Markdown 报告，包含：\n"
            "1) 职业画像与动机分析（3-6 条要点）\n"
            "2) 与目标职业的整体匹配度评级（例如：高度匹配/基本匹配/需谨慎）\n"
            "3) 关键优势/潜在风险点（各 3-5 条，结合答题内容给证据）\n"
            "4) 若坚持该职业的 4 周行动建议（按周分解）\n"
            "5) 若不适合该职业，建议的备选职业方向（至少 3 个，并解释理由）\n"
            "要求：只输出 Markdown，不要输出 JSON。"
        )

        user_prompt = (
            f"目标职业：{target_career}\n\n"
            "以下是用户的作答（字典形式，key 为题号，value 为选项文本）：\n"
            f"{json.dumps(req.answers, ensure_ascii=False, indent=2)}\n"
            "请围绕此目标职业，生成一份匹配度分析报告。"
        )

        markdown = _deepseek_markdown(system_prompt, user_prompt)
        return {"success": True, "markdown": markdown}


def _quantize_score(value: str, dimension_name: str) -> int:
    """
    将用户输入转换为 0-100 的分数
    
    参数:
        value: 用户输入（可能是数字字符串或文字描述）
        dimension_name: 维度名称（用于 AI 量化时的上下文）
    
    返回:
        0-100 的整数分数
    """
    value = str(value).strip()
    
    # 尝试直接解析为数字
    try:
        num = float(value)
        # 如果是 0-4 范围，转换为 0-100（假设 4 分制）
        if 0 <= num <= 4:
            return int((num / 4) * 100)
        # 如果是 0-100 范围，直接返回
        elif 0 <= num <= 100:
            return int(num)
        # 超出范围，限制在 0-100
        else:
            return max(0, min(100, int(num)))
    except ValueError:
        # 不是纯数字，需要调用 AI 量化
        pass
    
    # 调用 DeepSeek API 进行文字量化
    try:
        system_prompt = (
            "你是一位专业的竞争力评估专家。\n"
            "请根据用户提供的文字描述，将其量化为 0-100 的分数。\n"
            "评分标准：\n"
            "- 0-20：较差/无\n"
            "- 21-40：一般/较少\n"
            "- 41-60：中等/有一些\n"
            "- 61-80：良好/较多\n"
            "- 81-100：优秀/很多\n"
            "必须严格按照以下 JSON 格式返回，不要任何多余话术：\n"
            "{\n"
            '  "score": 数字（0-100）\n'
            "}"
        )
        
        user_prompt = (
            f"维度：{dimension_name}\n"
            f"用户输入：{value}\n\n"
            "请根据上述描述，给出 0-100 的量化分数。"
        )
        
        result = _deepseek_json(system_prompt, user_prompt)
        score = result.get("score", 50)  # 默认 50 分
        
        # 确保分数在 0-100 范围内
        score = max(0, min(100, int(score)))
        print(f"✅ [quantize_score] {dimension_name}: '{value}' → {score} 分")
        return score
    
    except Exception as e:
        print(f"❌ [quantize_score] AI 量化失败 ({dimension_name}: '{value}'): {e}")
        # AI 量化失败，返回默认分数 50
        return 50


@app.post("/api/analyze_natural_language")
def analyze_natural_language(req: AnalyzeNaturalLanguageRequest):
    """
    竞争力沙盘自然语言量化接口

    接收一段包含 GPA / 项目 / 实习 / 竞赛 / 英语 / 领导力 等信息的自然语言文本，
    调用 DeepSeek 提取并量化为 0-100 分数。

    返回示例：
    {
      "success": true,
      "scores": {
        "gpa": 95,
        "project": 80,
        "intern": 90,
        "competition": 100,
        "english": 90,
        "leader": 85
      },
      "fallback": false
    }
    """
    import traceback

    text = (req.text or "").strip()
    if not text:
        raise HTTPException(status_code=400, detail="text 不能为空")

    print(f"✅ [analyze_natural_language] 收到文本: {text[:80]}...")

    # 默认兜底分数（若 AI 失败则使用）
    fallback_scores = {
        "gpa": 50,
        "project": 50,
        "intern": 50,
        "competition": 50,
        "english": 50,
        "leader": 50,
    }

    try:
        system_prompt = (
            "你是一位专业的职业竞争力评估专家，擅长从自然语言描述中提取关键信息并量化。\n"
            "用户会给出一段关于自己竞争力的综合描述，可能包含：GPA/绩点、项目数量与难度、名企实习经历、竞赛获奖、英语水平、领导力与协作等。\n"
            "请你阅读用户的完整描述，根据以下规则，将 6 个维度量化为 0-100 分：\n"
            "- gpa：学业成绩/GPA/排名/奖学金等\n"
            "- project：项目实战（项目数量、复杂度、是否落地、是否与目标岗位相关）\n"
            "- intern：名企实习经历（公司级别、实习时长、岗位匹配度）\n"
            "- competition：竞赛获奖情况（校级/省级/国家级/国际级等）\n"
            "- english：英语学术能力（四六级/雅思托福/论文/报告等）\n"
            "- leader：领导力与协作（学生干部、项目负责人、团队协作经历等）\n"
            "输出格式：严格按照以下 JSON 结构返回，不要任何多余话术：\n"
            "{\n"
            '  \"scores\": {\n'
            '    \"gpa\": 0-100,\n'
            '    \"project\": 0-100,\n'
            '    \"intern\": 0-100,\n'
            '    \"competition\": 0-100,\n'
            '    \"english\": 0-100,\n'
            '    \"leader\": 0-100\n'
            "  }\n"
            "}\n"
            "注意：如果描述中某个维度完全缺失，请根据常识给出一个“保守中间值”（例如 40-60）而不是 0。"
        )

        user_prompt = f"以下是用户的自然语言描述：\n{text}"

        data = _deepseek_json(system_prompt, user_prompt)
        scores = data.get("scores") or {}

        def _norm(key: str, default: int) -> int:
            try:
                v = scores.get(key, default)
                v = int(float(v))
                return max(0, min(100, v))
            except Exception:
                return default

        normalized = {
            "gpa": _norm("gpa", fallback_scores["gpa"]),
            "project": _norm("project", fallback_scores["project"]),
            "intern": _norm("intern", fallback_scores["intern"]),
            "competition": _norm("competition", fallback_scores["competition"]),
            "english": _norm("english", fallback_scores["english"]),
            "leader": _norm("leader", fallback_scores["leader"]),
        }

        print(f"✅ [analyze_natural_language] 量化结果: {normalized}")

        return {
            "success": True,
            "scores": normalized,
            "fallback": False,
        }

    except HTTPException:
        raise
    except Exception as e:
        print(f"❌ [analyze_natural_language] 量化失败: {e}")
        print(f"❌ [analyze_natural_language] 错误堆栈: {traceback.format_exc()}")
        # AI 故障时使用兜底规则，保证前端功能可用
        return {
            "success": True,
            "scores": fallback_scores,
            "fallback": True,
        }


@app.post("/api/analyze_competitiveness")
def analyze_competitiveness(req: AnalyzeCompetitivenessRequest):
    """
    竞争力分析接口
    
    接收 6 个维度的用户输入（可以是数字或文字），进行量化后生成 AI 分析报告
    
    返回格式：
    {
      "success": true,
      "quantized_scores": {
        "gpa": 85,
        "project_experience": 70,
        "internship": 60,
        "competition": 75,
        "english_academic": 80,
        "leadership": 65
      },
      "analysis_report": "AI生成的Markdown格式分析报告",
      "fallback": false
    }
    """
    import traceback
    
    print(f"✅ [analyze_competitiveness] 收到竞争力分析请求")
    
    try:
        # 1. 量化所有维度分数
        print(f"🔄 [analyze_competitiveness] 开始量化各维度分数")
        
        quantized_scores = {
            "gpa": _quantize_score(req.gpa, "GPA"),
            "project_experience": _quantize_score(req.project_experience, "项目实战经验"),
            "internship": _quantize_score(req.internship, "名企实习经历"),
            "competition": _quantize_score(req.competition, "竞赛获奖情况"),
            "english_academic": _quantize_score(req.english_academic, "英语学术能力"),
            "leadership": _quantize_score(req.leadership, "领导力与协作"),
        }
        
        print(f"✅ [analyze_competitiveness] 量化完成: {quantized_scores}")
        
        # 2. 生成 AI 分析报告
        print(f"🔄 [analyze_competitiveness] 开始生成 AI 分析报告")
        
        system_prompt = (
            "你是一位专业的职业竞争力评估专家和生涯规划顾问。\n"
            "请根据用户提供的 6 个维度量化分数，生成一份个性化的竞争力分析报告。\n"
            "报告必须严格按照以下 Markdown 结构输出，不要任何多余话术：\n"
            "\n"
            "## 📊 竞争力总览\n"
            "（总体评分和一句话总结）\n"
            "\n"
            "## 📈 各维度详细分析\n"
            "### 1. GPA 学术成绩\n"
            "（分数：XX/100，评价和建议）\n"
            "\n"
            "### 2. 项目实战经验\n"
            "（分数：XX/100，评价和建议）\n"
            "\n"
            "### 3. 名企实习经历\n"
            "（分数：XX/100，评价和建议）\n"
            "\n"
            "### 4. 竞赛获奖情况\n"
            "（分数：XX/100，评价和建议）\n"
            "\n"
            "### 5. 英语学术能力\n"
            "（分数：XX/100，评价和建议）\n"
            "\n"
            "### 6. 领导力与协作\n"
            "（分数：XX/100，评价和建议）\n"
            "\n"
            "## 🎯 核心竞争力\n"
            "（列出 2-3 个最强维度）\n"
            "\n"
            "## ⚠️ 待提升领域\n"
            "（列出 2-3 个需要重点提升的维度）\n"
            "\n"
            "## 💡 个性化提升建议\n"
            "（针对待提升领域，给出 3-5 条具体可执行的建议）\n"
        )
        
        user_prompt = (
            "以下是用户在 6 个维度的量化分数：\n"
            f"- GPA 学术成绩：{quantized_scores['gpa']}/100\n"
            f"- 项目实战经验：{quantized_scores['project_experience']}/100\n"
            f"- 名企实习经历：{quantized_scores['internship']}/100\n"
            f"- 竞赛获奖情况：{quantized_scores['competition']}/100\n"
            f"- 英语学术能力：{quantized_scores['english_academic']}/100\n"
            f"- 领导力与协作：{quantized_scores['leadership']}/100\n\n"
            "请基于以上分数，生成一份详细的竞争力分析报告。"
        )
        
        analysis_report = _deepseek_markdown(system_prompt, user_prompt)
        
        print(f"✅ [analyze_competitiveness] AI 分析报告生成成功，长度: {len(analysis_report)} 字符")
        
        return {
            "success": True,
            "quantized_scores": quantized_scores,
            "analysis_report": analysis_report,
            "fallback": False
        }
    
    except Exception as e:
        print(f"❌ [analyze_competitiveness] 生成失败: {e}")
        print(f"❌ [analyze_competitiveness] 错误堆栈: {traceback.format_exc()}")
        
        # 降级逻辑：返回基础分析
        quantized_scores_fallback = {
            "gpa": _quantize_score(req.gpa, "GPA") if hasattr(req, 'gpa') else 0,
            "project_experience": _quantize_score(req.project_experience, "项目实战经验") if hasattr(req, 'project_experience') else 0,
            "internship": _quantize_score(req.internship, "名企实习经历") if hasattr(req, 'internship') else 0,
            "competition": _quantize_score(req.competition, "竞赛获奖情况") if hasattr(req, 'competition') else 0,
            "english_academic": _quantize_score(req.english_academic, "英语学术能力") if hasattr(req, 'english_academic') else 0,
            "leadership": _quantize_score(req.leadership, "领导力与协作") if hasattr(req, 'leadership') else 0,
        }
        
        fallback_report = (
            "## 📊 竞争力总览\n\n"
            "基于您提供的 6 个维度数据，系统已进行初步分析。\n\n"
            "## 📈 各维度分数\n\n"
            f"- GPA 学术成绩：{quantized_scores_fallback['gpa']}/100\n"
            f"- 项目实战经验：{quantized_scores_fallback['project_experience']}/100\n"
            f"- 名企实习经历：{quantized_scores_fallback['internship']}/100\n"
            f"- 竞赛获奖情况：{quantized_scores_fallback['competition']}/100\n"
            f"- 英语学术能力：{quantized_scores_fallback['english_academic']}/100\n"
            f"- 领导力与协作：{quantized_scores_fallback['leadership']}/100\n\n"
            "## 💡 建议\n\n"
            "建议重点关注分数较低的维度，制定针对性的提升计划。"
        )
        
        return {
            "success": True,
            "quantized_scores": quantized_scores_fallback,
            "analysis_report": fallback_report,
            "fallback": True
        }


@app.post("/api/generate-career")
def generate_career(req: GenerateCareerRequest):
    """
    接收：性格测试 JSON + 虚拟实验 Markdown + 可选补充说明
    输出：整合后的生涯规划 Markdown 报告
    """
    system_prompt = (
        "你是一位资深生涯规划师。你将整合两份输入：\n"
        "- 性格测试结果（JSON：可能含截图/自述/字段）\n"
        "- 虚拟实验倾向分析（Markdown）\n"
        "请输出一份最终的生涯规划 Markdown 报告，包含：\n"
        "1) 个人画像（性格/动机/工作方式偏好）\n"
        "2) 目标职业方向建议（3 个主方向 + 3 个备选方向）\n"
        "3) 方向匹配理由（用证据对齐：来自性格测试与虚拟实验）\n"
        "4) 能力差距清单（按：基础/项目/软技能/行业认知）\n"
        "5) 12 周成长路线图（按周分解，每周 3-6 个任务）\n"
        "6) 作品集/项目建议（至少 3 个可落地项目，写清楚产出物）\n"
        "7) 简历与面试策略（关键词、故事线、STAR/项目讲法）\n"
        "要求：只输出 Markdown，不要输出 JSON。"
    )

    user_prompt = (
        "【性格测试 JSON】\n"
        f"{json.dumps(req.personality_json, ensure_ascii=False, indent=2)}\n\n"
        "【虚拟实验 Markdown】\n"
        f"{req.experiment_markdown}\n\n"
        "【用户补充说明（可为空）】\n"
        f"{req.note or ''}\n\n"
        "请输出最终的生涯规划 Markdown 报告。"
    )

    markdown = _deepseek_markdown(system_prompt, user_prompt)
    return {"success": True, "markdown": markdown}


@app.post("/api/generate-interview-report")
def generate_interview_report(req: GenerateInterviewReportRequest):
    """
    生成面试分析报告
    接收：完整的对话历史记录和目标岗位
    输出：Markdown 格式的面试分析报告
    """
    system_prompt = (
        "你是一位长期为大学生（本科生 + 研究生）做模拟面试辅导的资深面试官兼职业发展顾问。\n"
        "你会基于【系统已提取的元信息】和【完整对话记录】生成一份**高度贴合本次面试表现的个性化分析报告**，而不是模板化的空洞总结。\n\n"
        "【总体要求】\n"
        "- 报告对象限定为大学生：需要区分本科生 / 研究生的典型特点和面试关注点；\n"
        "- 所有评价、建议必须紧密结合本次对话中的**具体回答内容**，避免使用任何“通用模板化”的空话；\n"
        "- 采用 Markdown 输出，结构清晰、层级明确、便于直接作为报告下载保存；\n"
        "- 语言要专业、客观、友好，适合大学生阅读，避免居高临下或过度严厉的语气。\n\n"
        "【报告结构（必须包含以下 7 个模块）】\n\n"
        "1）报告基本信息\n"
        "- 使用系统提供的元信息（身份、方向、时长、生成时间等），整理成一个简要信息区块；\n"
        "- 不要自行编造与元信息相矛盾的内容，可以在元信息缺失时给出“未明确”或“未在对话中提及”的描述。\n\n"
        "2）面试整体评价\n"
        "- 结合对话中用户的整体表现，给出一段 2–3 段落的综合评价；\n"
        "- 需要同时覆盖：专业基础、表达与沟通、结构化思维、临场应变、与岗位/方向的匹配度等；\n"
        "- 评价中要引用少量对话中的**具体回答片段或现象**作为依据，而不是只给抽象形容词。\n\n"
        "3）综合评分（百分制 + 维度评分）\n"
        "- 给出一个总分（0–100 分），并给出 3–5 个评分维度（如“专业基础”“项目/实践”“表达与沟通”“逻辑与结构”“岗位匹配度”）；\n"
        "- 每个维度给出分数，并用 1–2 句话解释打分理由，理由要引用对话中的真实表现（例如“在 XX 问题中能回答出关键概念，但在 XX 追问时显得不够系统”）。\n\n"
        "4）答题亮点总结\n"
        "- 提炼 1–3 个**真实存在于本次对话中的亮点**，例如：\n"
        "  - 在某个专业问题上的回答体现出扎实的课程基础；\n"
        "  - 在介绍项目/实践经历时条理清晰、能够量化结果；\n"
        "  - 在自我介绍或职业动机上有明确的自我认知；\n"
        "- 每个亮点需要简要指出“体现在哪类问题中”“为什么是亮点”。\n\n"
        "5）主要提升方向\n"
        "- 针对本次面试中暴露出的具体问题给出 2–4 条改进建议；\n"
        "- 每条建议都要对应到某一类问题或具体表现，例如“技术细节题回答模糊”“表达缺少结构”“对目标岗位缺少了解”等；\n"
        "- 建议要可落地，如“建议针对 XX 主题整理 3–5 个高频题并写出演练稿”“建议使用 STAR 结构重写某个项目经历”。\n\n"
        "6）逐题分析（核心模块）\n"
        "- 针对本次面试中出现的每一个关键问题（可以按 3–8 个代表性问题汇总），输出一个子模块：\n"
        "  - 问题：简要概括该轮提问的核心问题；\n"
        "  - 用户回答评价：结合用户原始回答，指出优点和需要改进的地方（引用回答中的关键信息）；\n"
        "  - 标准参考答案：给出贴合大学生/目标岗位的专业参考回答，可以用结构化的小标题或分点形式；\n"
        "- 参考答案要保证：逻辑清晰、信息充分、技术或专业内容大体正确、用语简洁易懂。\n\n"
        "7）总结与后续建议\n"
        "- 结合用户当前身份（本科生/研究生，如元信息中未给出则用“在校大学生”表述），给出面试准备与职业发展的建议；\n"
        "- 建议可以覆盖：短期（1–3 个月）的面试准备、简历与项目打磨、课程/科研侧重方向等；\n"
        "- 保持鼓励性与建设性，避免只指出问题不提供可行路径。\n\n"
        "【重要限制】\n"
        "- 不要简单照搬“通用模板”的话术，要根据本次对话的具体内容做针对性的描述；\n"
        "- 不要虚构用户并未在对话中体现的经历（如不存在的实习/项目/科研），可以对已有经历做合理的总结与提炼；\n"
        "- 报告中可以适度进行信息归纳和推断，但不要杜撰事实。\n"
    )
    
    # 构建对话历史文本
    chat_text = ""
    for msg in req.chat_history:
        role = "面试官" if msg.get("role") == "ai" else "求职者"
        content = msg.get("content", "")
        chat_text += f"{role}：{content}\n\n"

    meta = req.meta or {}
    direction = meta.get("direction") or "未在对话中明确"
    identity = meta.get("identity") or "未在对话中明确"
    duration = meta.get("durationMinutes")
    duration_text = f"{duration} 分钟" if isinstance(duration, (int, float)) and duration > 0 else "未统计"
    generated_at = meta.get("generatedAt") or "系统时间未提供"

    user_prompt = (
        "【元信息（已由系统预提取）】\n"
        f"- 面试方向：{direction}\n"
        f"- 面试者身份：{identity}\n"
        f"- 面试时长：{duration_text}\n"
        f"- 报告生成时间：{generated_at}\n"
        f"- 目标岗位（如有）：{req.target_role or '未指定'}\n\n"
        "【完整对话记录】（按照时间顺序，从开场到结束）：\n\n"
        f"{chat_text}\n"
        "请严格依据以上元信息和对话内容，按照系统提示的 7 个模块结构生成一份面向大学生的个性化面试分析报告。"
    )
    
    try:
        markdown = _deepseek_markdown(system_prompt, user_prompt)
        return {"success": True, "markdown": markdown}
    except Exception as e:
        return {"success": False, "error": str(e)}

# ==========================================
#  Production: Serve Vue frontend
# ==========================================
@app.get("/{full_path:path}")
async def serve_frontend(full_path: str):
    """Serve Vue frontend for non-API routes in production"""
    frontend_index = os.path.join(os.path.dirname(__file__), "..", "frontend", "dist", "index.html")
    if os.path.exists(frontend_index):
        return FileResponse(frontend_index)
    return {"message": "Frontend not built. Run 'npm run build' in frontend directory."}

# ==========================================
#  启动入口
# ==========================================
if __name__ == "__main__":
    # ==========================================
    #  数据库连接测试（可选，用于验证配置）
    # ==========================================
    print("=" * 50)
    print("📊 数据库连接测试")
    print("=" * 50)
    
    # 测试1: 数据库连接
    print("\n1️⃣ 测试数据库连接...")
    try:
        from .db_config import get_db_cursor
        conn, cursor = get_db_cursor()
        cursor.close()
        conn.close()
        print("✅ 数据库连接成功！")
    except Exception as e:
        print(f"❌ 数据库连接失败：{e}")
        print("   提示：请检查 Render 环境变量配置（DB_HOST, DB_PORT, DB_USER, DB_PASSWORD, DB_NAME）")
    
    # 测试2: 获取所有用户
    print("\n2️⃣ 测试获取所有用户数据...")
    all_users = get_all_users()
    print(f"✅ 成功获取 {len(all_users)} 条用户数据")
    if len(all_users) > 0:
        print(f"   示例用户：{all_users[0].get('username', 'N/A')}")
    
    # 测试3: 用户登录验证
    print("\n3️⃣ 测试用户登录验证...")
    if len(all_users) > 0:
        test_user = all_users[0]
        test_username = test_user.get('username', '')
        test_password = test_user.get('password', '')
        
        # 测试正确密码
        success, msg = user_login(test_username, test_password)
        print(f"   正确密码测试: {msg}")
        
        # 测试错误密码
        success, msg = user_login(test_username, "wrong_password")
        print(f"   错误密码测试: {msg}")
    else:
        print("   ⚠️ 无用户数据，跳过登录测试")
    
    print("\n" + "=" * 50)
    print("✅ 测试完成！")
    print("=" * 50)
    print("\n🚀 启动 FastAPI 服务器...")
    print("   访问地址: http://127.0.0.1:8001")
    print("   API 文档: http://127.0.0.1:8001/docs\n")
    
    uvicorn.run(app, host="127.0.0.1", port=8001)

if __name__ == "__main__":
    uvicorn.run(app, host="0.0.0.0", port=8000)